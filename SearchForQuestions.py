from logging import error
from os import write
from typing import runtime_checkable
import json
import io
import random

import docx
import docx
from docx.shared import RGBColor


def get_json(doc,item):
    try:
        doc = docx.Document(f"{doc}")
        table = doc.tables[0]
        data = []
        result = []
        keys = None
        for table in doc.tables:
            table
            for i, row in enumerate(table.rows):
                text = (cell.text for cell in row.cells)
                if i == 0:
                    keys = tuple(text)
                    continue
                row_data = dict(zip(keys, text))
                data.append(row_data)
                
        for table in doc.tables:
            for row in table.rows:
                j = 0
                for cell in row.cells:
                    j+= 1
                    if len(data[1]) == 2:
                        if j == 1:
                            array = []
                            array.append(cell.text)
                    elif len(data[1]) == 3:
                        if j == 1:
                            array = []
                            array.append(cell.text)
                        elif j == 2:
                            array.append(cell.text)
                    for para in cell.paragraphs:
                        for run in para.runs:
                            check = RGBColor.from_string('FF0000')
                            if run.font.color.rgb == check :
                                if run.text != " " and run.text != "\n":
                                    text1 = para.text
                                    text_array = []
                                    if len(data[1]) == 2:
                                        text_array.append(array[0])
                                        text_array.append(text1)
                                    elif len(data[1]) == 3:
                                        text_array.append(array[0])
                                        text_array.append(array[1])
                                        text_array.append(text1)
                                    row_data = dict(zip(keys, text_array))
                                    result.append(row_data)

        seen = set()
        new_result = []
        for d in result:
            t = tuple(sorted(d.items()))
            if t not in seen:
                seen.add(t)
                new_result.append(d)
        for d in data:
            for n in new_result:
                if list(n.items())[0][1] == list(d.items())[0][1]:
                    d["Ответ"] = list(n.items())[1][1]

        with io.open(f"TestJson/{item.faculty}/{item.semester}/{item.year}/{item.group}/table{item.document}.json", "w", encoding="utf-8") as data_file:
            json.dump(data, data_file, indent=4, ensure_ascii=False)
    except:
        print(error)

def shuf(item):
    try:
        with io.open(f"TestJson/{item.faculty}/{item.semester}/{item.year}/{item.group}/table{item.document}.json", encoding="utf-8") as f:
            templates = json.load(f)
        random.shuffle(templates)
        with io.open(f"TestJson/{item.faculty}/{item.semester}/{item.year}/{item.group}/table{item.document}shuf.json", "w", encoding="utf-8") as data_file:
            json.dump(templates, data_file, indent=4, ensure_ascii=False)
    except:
        print(error)
